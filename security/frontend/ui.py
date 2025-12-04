from __future__ import annotations

import io
import re
import time
from pathlib import Path
from typing import Optional

import pytesseract
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse  # ✨ added RedirectResponse
from PIL import Image, UnidentifiedImageError, ImageOps, ImageFilter, ImageStat
from pytesseract import TesseractNotFoundError
from fastapi.staticfiles import StaticFiles

try:
    import fitz  # PyMuPDF for PDFs
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

# -------------------- Import modules --------------------
# Ensure package imports resolve when embedded under backend-api
from security.reports.report_service import generate_pdf
from security.strategies import load_strategies, ALLOWED_STRATEGIES, get_checker

# ✨ added
from collections import deque
from datetime import datetime

# -------------------- environment checks --------------------
_HAS_TESSERACT = True
_TESSERACT_VERSION = ""
try:
    _TESSERACT_VERSION = str(pytesseract.get_tesseract_version())
except Exception:
    _HAS_TESSERACT = False


# -------------------- error helpers --------------------
def error_response(status_code: int, code: str, message: str, errors: Optional[list] = None):
    errs = errors or [{"code": code, "message": message}]
    payload = {
        "ok": False,
        "error": True,
        "has_errors": len(errs) > 0,
        "code": code,
        "detail": message,
        "errors": errs,
    }
    return JSONResponse(payload, status_code=status_code)


# -------------------- utility --------------------
def _safe_uid(s: str) -> str:
    """Make a filesystem-safe id (no spaces/odd chars)."""
    s = re.sub(r"\s+", "_", s.strip())
    return re.sub(r"[^A-Za-z0-9._-]", "-", s)


# -------------------- extraction helpers --------------------
def _ocr_image_bytes(data: bytes) -> str:
    """OCR bytes of an image; if Tesseract is missing, return empty text instead of 500."""
    try:
        with Image.open(io.BytesIO(data)) as img:
            try:
                # Normalize dark-theme screenshots for better OCR.
                g = img.convert("L")
                mean_luma = ImageStat.Stat(g).mean[0] if ImageStat else 128
                if mean_luma < 80:
                    g = ImageOps.invert(g)
                g = ImageOps.autocontrast(g)
                g = g.filter(ImageFilter.SHARPEN)
                return pytesseract.image_to_string(g)
            except TesseractNotFoundError:
                return ""  # no OCR available → behave gracefully
    except (UnidentifiedImageError, OSError):
        return ""


def _extract_pdf_bytes(
    data: bytes, previews_dir: Path, stem: str
) -> tuple[str, Optional[Path]]:
    print("[extract_pdf_bytes] start", {"bytes": len(data), "stem": stem})
    if not fitz:
        return "", None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        text = "".join((page.get_text() or "") for page in doc)
        preview_path: Optional[Path] = None
        if len(doc) > 0:
            pix = doc[0].get_pixmap()
            previews_dir.mkdir(parents=True, exist_ok=True)
            preview_path = previews_dir / f"{stem}_page1.png"
            pix.save(str(preview_path))
        doc.close()
        print("[extract_pdf_bytes] done", {"chars": len(text), "preview": str(preview_path) if preview_path else None})
        return text, preview_path
    except Exception:
        return "", None


def _extract_docx_bytes(data: bytes, previews_dir: Path) -> tuple[str, Optional[Path]]:
    """
    Extract text from DOCX, and OCR any embedded images as a fallback.
    Returns (text, preview_path).
    """
    if not DocxDocument:
        return "", None
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        ocr_parts: list[str] = []
        preview_path: Optional[Path] = None
        for rel in doc.part._rels.values():
            # Only process image relationships
            reltype = getattr(rel, "reltype", "") or ""
            if "image" not in reltype:
                continue
            target = getattr(rel, "target_part", None)
            blob = getattr(target, "blob", None)
            if not blob:
                continue
            # OCR the embedded image
            ocr_text = _ocr_image_bytes(blob)
            if ocr_text.strip():
                ocr_parts.append(ocr_text)
            # Save first image as preview
            if not preview_path:
                previews_dir.mkdir(parents=True, exist_ok=True)
                # rel.target_ref looks like 'media/image1.png'
                fname = Path(getattr(target, "partname", "image"))
                preview_path = previews_dir / _safe_uid(fname.name)
                with open(preview_path, "wb") as f:
                    f.write(blob)

        full_text = "\n".join(t for t in ["\n".join(parts)] + ocr_parts if t)
        return full_text, preview_path
    except Exception:
        return "", None


def extract_text_and_preview_bytes(
    filename: str, data: bytes, previews_dir: Path
) -> tuple[str, Optional[Path]]:
    ext = Path(filename).suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        text = _ocr_image_bytes(data)
        previews_dir.mkdir(parents=True, exist_ok=True)
        preview_path = previews_dir / _safe_uid(filename)
        with open(preview_path, "wb") as f:
            f.write(data)
        return text, preview_path
    if ext == ".pdf":
        return _extract_pdf_bytes(data, previews_dir, Path(_safe_uid(filename)).stem)
    if ext == ".docx":
        return _extract_docx_bytes(data, previews_dir)
    if ext in {
        ".txt",
        ".log",
        ".reg",
        ".csv",
        ".ini",
        ".json",
        ".xml",
        ".htm",
        ".html",
    }:
        return data.decode("utf-8", errors="ignore"), None
    return "", None


# -------------------- FastAPI app --------------------
app = FastAPI(title="AutoAudit Evidence Scanner")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve files in frontend 
FRONTEND_DIR = Path(__file__).resolve().parent
app.mount("/frontend", StaticFiles(directory=FRONTEND_DIR), name="frontend")

# Paths relative to security/
ROOT = Path(__file__).resolve().parents[1]   # .../security
RESULTS = ROOT / "results"                   # security/results/...
TEMPLATES = RESULTS                          # report_template.docx lives here
OUT_DIR = RESULTS / "reports"                # PDF (or DOCX/TXT) outputs
PREVIEWS = RESULTS / "previews"              # preview images for evidence
INDEX_HTML = ROOT / "frontend" / "ui.html"      # serve the UI from frontend/ui.html

# ✨ Recent scan: in-memory log
SCAN_MEM = deque(maxlen=50)

def _push_mem_log(user: str, strategy: str, status: str) -> None:
    """status: 'success' | 'error'"""
    SCAN_MEM.appendleft({
        "ts": datetime.now().astimezone().isoformat(),
        "user": user,
        "strategy": strategy,
        "status": status,
    })


@app.get("/", response_class=HTMLResponse)
def index():
    if INDEX_HTML.exists():
        return INDEX_HTML.read_text(encoding="utf-8", errors="ignore")
    return "<h3>AutoAudit Evidence Scanner</h3>"


@app.get("/strategies")
def api_strategies():
    """Return only the curated benchmark strategies expected by the UI."""
    try:
        # preserve canonical order from ALLOWED_STRATEGIES
        return [
            {"name": name, "description": desc}
            for name, desc in ALLOWED_STRATEGIES.items()
        ]
    except Exception as exc:
        return error_response(
            status_code=500,
            code="STRATEGY_LOAD_FAILED",
            message="Unable to load strategies.",
            errors=[{"code": "STRATEGY_LOAD_FAILED", "message": str(exc)}],
        )


@app.get("/health")
def health():
    """Quick status to debug issues without crashing the UI."""
    try:
        return {
            "ok": True,
            "has_tesseract": _HAS_TESSERACT,
            "tesseract_version": _TESSERACT_VERSION,
            "template_exists": (TEMPLATES / "report_template.docx").exists(),
            "previews_dir": str(PREVIEWS),
            "out_dir": str(OUT_DIR),
            "error": False,
            "errors": [],
        }
    except Exception as exc:
        return error_response(
            status_code=500,
            code="HEALTH_CHECK_FAILED",
            message="Health check failed.",
            errors=[{"code": "HEALTH_CHECK_FAILED", "message": str(exc)}],
        )

# ✨ Recent scan: APIs + page + aliases
@app.post("/api/scan-mem-log")
def api_post_scan_mem_log(payload: dict):
    _push_mem_log(
        user=str(payload.get("user") or payload.get("user_id") or "user"),
        strategy=str(payload.get("strategy") or payload.get("strategy_name") or ""),
        status=str(payload.get("status") or "success"),
    )
    return {"ok": True}

@app.get("/api/scan-mem-log")
def api_get_scan_mem_log():
    return JSONResponse(list(SCAN_MEM))

@app.get("/scan-mem", response_class=HTMLResponse)
def scan_mem_page():
    rows = "".join(
        f"<tr><td>{r.get('ts','')}</td>"
        f"<td>{r.get('user','')}</td>"
        f"<td>{r.get('strategy','')}</td>"
        f"<td>{r.get('status','')}</td></tr>"
        for r in SCAN_MEM
    ) or "<tr><td colspan='4'>No runs yet</td></tr>"
    return HTMLResponse(
        f"""<!doctype html>
<meta charset="utf-8">
<title>Recent Scans — AutoAudit</title>
<style>
  body{{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;background:#fff;color:#111;margin:24px}}
  a{{color:#0d6efd;text-decoration:none}}
  table{{border-collapse:collapse;width:100%;max-width:1040px;margin-top:12px}}
  th,td{{border:1px solid #e5e7eb;padding:10px;font-size:14px}}
  th{{background:#f9fafb;text-align:left}}
  .ok{{color:#047857;font-weight:700}}
  .bad{{color:#b00020;font-weight:700}}
</style>
<h1 style="margin:0 0 8px 0;">Recent Scans</h1>
<p><a href="/">← Back to scanner</a></p>
<table>
  <tr><th>Time</th><th>User</th><th>Strategy</th><th>Status</th></tr>
  {rows}
</table>"""
    )

@app.get("/recent-scans", include_in_schema=False)
def recent_scans_redirect():
    return RedirectResponse(url="/scan-mem")

@app.get("/scan-log", include_in_schema=False)
def scan_log_redirect():
    return RedirectResponse(url="/scan-mem")


@app.post("/scan")
async def scan(
    evidence: UploadFile = File(...),
    strategy_name: str = Form(...),  # preserved for frontend compatibility
    strategy: str | None = Form(None),  # alternative field name
    user_id: str = Form("user"),
):
    try:
        strategy_label = strategy or strategy_name
        strat_obj = get_checker(strategy_label) if 'get_checker' in globals() else None
        if strat_obj is None:
            # fallback to previous lookup
            strat_obj = next((s for s in load_strategies() if s.name == strategy_label), None)

        if not strat_obj:
            _push_mem_log(user_id, strategy_label, "error")
            return error_response(
                status_code=400,
                code="STRATEGY_NOT_FOUND",
                message=f"Unknown strategy: {strategy_label}",
            )

        # read upload
        content = await evidence.read()
        if not content:
            _push_mem_log(user_id, strat_obj.name, "error")
            return error_response(
                status_code=400,
                code="EMPTY_UPLOAD",
                message="Evidence file is empty.",
            )

        # extract text + preview
        text, preview_path = extract_text_and_preview_bytes(
            evidence.filename, content, PREVIEWS
        )
        if not text.strip():
            _push_mem_log(user_id, strat_obj.name, "success")  # log a successful run even without readable text
            return JSONResponse(
                {
                    "ok": True,
                    "error": False,
                    "has_errors": False,
                    "errors": [],
                    "findings": [],
                    "reports": [],
                    "note": "No readable text found in evidence."
                }
            )

        # run rules
        if hasattr(strat_obj, "emit_hits"):
            try:
                findings = strat_obj.emit_hits(text, source_file=evidence.filename, user_id=user_id) or []
            except TypeError:
                # fallback for strategies that don't accept extra kwargs
                findings = strat_obj.emit_hits(text) or []
        else:
            hits = strategy.match(text) or []
            findings = (
                [{
                    "test_id": "",
                    "sub_strategy": "",
                    "detected_level": "",
                    "pass_fail": "",
                    "priority": "",
                    "recommendation": "",
                    "evidence": hits,
                }]
                if hits else []
            )

        # Normalize findings to a common schema and ensure details are present
        def _clip(s: str, n: int = 400) -> str:
            if not s:
                return ""
            s = " ".join(str(s).split())
            return s if len(s) <= n else s[: n - 3] + "..."

        normalized = []
        for f in findings:
            f = dict(f)
            f.setdefault("test_id", "")
            f.setdefault("sub_strategy", "")
            f.setdefault("detected_level", "")
            f.setdefault("pass_fail", "UNKNOWN")
            f.setdefault("priority", "")
            f.setdefault("recommendation", "")
            f.setdefault("evidence", [])

            # coerce evidence to list[str]
            ev = f.get("evidence")
            if isinstance(ev, str):
                ev_list = [ev]
            elif isinstance(ev, list):
                ev_list = [str(x) for x in ev]
            else:
                ev_list = []
            f["evidence"] = ev_list

            # details (observed/expected)
            obs_full = ""
            if ev_list:
                obs_full = ev_list[0]
            exp_full = f.get("recommendation") or f.get("description") or ""

            det = f.get("details") or {}
            det.setdefault("observed_full", obs_full)
            det.setdefault("expected_full", exp_full)
            det.setdefault("observed", _clip(det.get("observed_full", "")))
            det.setdefault("expected", _clip(det.get("expected_full", "")))
            f["details"] = det

            normalized.append(f)

        findings = normalized

        # generate reports (PDF preferred; fallback DOCX/TXT so we never 500)
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        generated: list[str] = []
        note = ""
        errors: list[dict] = []

        for r in findings:
            base_uid = _safe_uid(f"{user_id}-{strat_obj.name}-{Path(evidence.filename).stem}")
            # make name unique to avoid 'Permission denied' if file is locked from previous run
            uid = f"{base_uid}-{int(time.time())}"

            # payload for the existing template flow
            payload = {
                "UniqueID": uid,
                "UserID": user_id,
                "Evidence": evidence.filename,
                "Evidence Preview": str(preview_path) if preview_path else "",
                "Strategy": strat_obj.name,
                "TestID": r.get("test_id", ""),
                "Sub-Strategy": r.get("sub_strategy", ""),
                "ML Level": r.get("detected_level", ""),
                "Pass/Fail": r.get("pass_fail", ""),
                "Priority": r.get("priority", ""),
                "Recommendation": r.get("recommendation", ""),
                "Evidence Extract": "; ".join(r.get("evidence", [])),
                "Description": r.get("description", ""),
                "Confidence": r.get("confidence", ""),
            }

            try:
                pdf_path = generate_pdf(
                    payload,
                    template_path=str(TEMPLATES / "report_template.docx"),
                    output_dir=str(OUT_DIR),
                    base_dir=str(ROOT),
                )
                generated.append(Path(pdf_path).name)
            except Exception as exc:
                # If report generation fails even with fallbacks, capture the error
                errors.append({"code": "REPORT_GENERATION_FAILED", "message": str(exc)})

        _push_mem_log(user_id, strat_obj.name, "success")  # ✨ log success
        return {
            "ok": True,
            "error": False,
            "has_errors": len(errors) > 0,
            "errors": errors,
            "findings": findings,
            "reports": generated,
            "note": note,
        }

    except HTTPException as exc:
        _push_mem_log(user_id if 'user_id' in locals() else "user", strategy_name, "error")  # ✨ log error
        return error_response(
            status_code=exc.status_code,
            code="VALIDATION_ERROR",
            message=str(exc.detail),
        )
    except Exception as exc:
        _push_mem_log(user_id if 'user_id' in locals() else "user", strategy_name, "error")  # ✨ log error
        return error_response(
            status_code=500,
            code="SERVER_ERROR",
            message="Unexpected error during scan.",
            errors=[{"code": "SERVER_ERROR", "message": str(exc)}],
        )


@app.get("/reports/{filename}")
def download_report(filename: str):
    path = OUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report not found")

    media = (
        "application/pdf"
        if filename.lower().endswith(".pdf")
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if filename.lower().endswith(".docx")
        else "text/plain"
    )
    return FileResponse(str(path), media_type=media, filename=filename)
