from __future__ import annotations

import io
import re
import time
from pathlib import Path
from typing import Optional

import pytesseract
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from PIL import Image, UnidentifiedImageError
from pytesseract import TesseractNotFoundError

try:
    import fitz  # PyMuPDF for PDFs
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:
    DocxDocument = None

from reports.report_service import generate_pdf
from strategies import load_strategies

# -------------------- environment checks --------------------
_HAS_TESSERACT = True
_TESSERACT_VERSION = ""
try:
    _TESSERACT_VERSION = str(pytesseract.get_tesseract_version())
except Exception:
    _HAS_TESSERACT = False


# -------------------- utility --------------------
def _safe_uid(s: str) -> str:
    """Make a filesystem-safe id (no spaces/odd chars)."""
    s = re.sub(r"\s+", "_", s.strip())
    return re.sub(r"[^A-Za-z0-9._-]", "-", s)


# -------------------- extraction helpers --------------------
def _ocr_image_bytes(data: bytes) -> str:
    """OCR bytes of an image; if Tesseract is missing, return empty text instead of 500."""
    try:
        with Image.open(io.BytesIO(data)) as img:
            try:
                return pytesseract.image_to_string(img)
            except TesseractNotFoundError:
                return ""  # no OCR available → behave gracefully
    except (UnidentifiedImageError, OSError):
        return ""


def _extract_pdf_bytes(
    data: bytes, previews_dir: Path, stem: str
) -> tuple[str, Optional[Path]]:
    if not fitz:
        return "", None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        text = "".join((page.get_text() or "") for page in doc)
        preview_path: Optional[Path] = None
        if len(doc) > 0:
            pix = doc[0].get_pixmap()
            previews_dir.mkdir(parents=True, exist_ok=True)
            preview_path = previews_dir / f"{stem}_page1.png"
            pix.save(str(preview_path))
        doc.close()
        return text, preview_path
    except Exception:
        return "", None


def _extract_docx_bytes(data: bytes) -> str:
    if not DocxDocument:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_and_preview_bytes(
    filename: str, data: bytes, previews_dir: Path
) -> tuple[str, Optional[Path]]:
    ext = Path(filename).suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        text = _ocr_image_bytes(data)
        previews_dir.mkdir(parents=True, exist_ok=True)
        preview_path = previews_dir / _safe_uid(filename)
        with open(preview_path, "wb") as f:
            f.write(data)
        return text, preview_path
    if ext == ".pdf":
        return _extract_pdf_bytes(data, previews_dir, Path(_safe_uid(filename)).stem)
    if ext == ".docx":
        return _extract_docx_bytes(data), None
    if ext in {
        ".txt",
        ".log",
        ".reg",
        ".csv",
        ".ini",
        ".json",
        ".xml",
        ".htm",
        ".html",
    }:
        return data.decode("utf-8", errors="ignore"), None
    return "", None


# -------------------- FastAPI app --------------------
app = FastAPI(title="AutoAudit Evidence Scanner")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Paths relative to security/
ROOT = Path(__file__).resolve().parents[1]   # .../security
RESULTS = ROOT / "results"                   # security/results/...
TEMPLATES = RESULTS                          # report_template.docx lives here
OUT_DIR = RESULTS / "reports"                # PDF (or DOCX/TXT) outputs
PREVIEWS = RESULTS / "previews"              # preview images for evidence
INDEX_HTML = ROOT / "aa_ui" / "ui.html"      # serve the UI from aa_ui/ui.html


@app.get("/", response_class=HTMLResponse)
def index():
    if INDEX_HTML.exists():
        return INDEX_HTML.read_text(encoding="utf-8", errors="ignore")
    return "<h3>AutoAudit Evidence Scanner</h3>"


@app.get("/strategies")
def api_strategies():
    out = []
    for s in load_strategies():
        try:
            desc = s.description()
        except Exception:
            desc = ""
        out.append({"name": s.name, "description": desc})
    return out


@app.get("/health")
def health():
    """Quick status to debug issues without crashing the UI."""
    return {
        "ok": True,
        "has_tesseract": _HAS_TESSERACT,
        "tesseract_version": _TESSERACT_VERSION,
        "template_exists": (TEMPLATES / "report_template.docx").exists(),
        "previews_dir": str(PREVIEWS),
        "out_dir": str(OUT_DIR),
    }


@app.post("/scan")
async def scan(
    evidence: UploadFile = File(...),
    strategy_name: str = Form(...),
    user_id: str = Form("user"),
):
    # find strategy
    strategy = next((s for s in load_strategies() if s.name == strategy_name), None)
    if not strategy:
        raise HTTPException(status_code=400, detail=f"Unknown strategy: {strategy_name}")

    # read upload
    content = await evidence.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty upload")

    # extract text + preview
    text, preview_path = extract_text_and_preview_bytes(
        evidence.filename, content, PREVIEWS
    )
    if not text.strip():
        return JSONResponse(
            {"ok": True, "findings": [], "reports": [], "note": "No readable text found in evidence."}
        )

    # run rules
    if hasattr(strategy, "emit_hits"):
        findings = strategy.emit_hits(text, source_file=evidence.filename) or []
    else:
        hits = strategy.match(text) or []
        findings = (
            [{
                "test_id": "",
                "sub_strategy": "",
                "detected_level": "",
                "pass_fail": "",
                "priority": "",
                "recommendation": "",
                "evidence": hits,
            }]
            if hits else []
        )

    # generate reports (PDF preferred; fallback DOCX/TXT so we never 500)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated: list[str] = []
    note = ""

    for r in findings:
        base_uid = _safe_uid(f"{user_id}-{strategy.name}-{Path(evidence.filename).stem}")
        # make name unique to avoid 'Permission denied' if file is locked from previous run
        uid = f"{base_uid}-{int(time.time())}"

        # payload for the existing template flow
        payload = {
            "UniqueID": uid,
            "UserID": user_id,
            "Evidence": evidence.filename,
            "Evidence Preview": str(preview_path) if preview_path else "",
            "Strategy": strategy.name,
            "TestID": r.get("test_id", ""),
            "Sub-Strategy": r.get("sub_strategy", ""),
            "ML Level": r.get("detected_level", ""),
            "Pass/Fail": r.get("pass_fail", ""),
            "Priority": r.get("priority", ""),
            "Recommendation": r.get("recommendation", ""),
            "Evidence Extract": "; ".join(r.get("evidence", [])),
            "Description": r.get("description", ""),
            "Confidence": r.get("confidence", ""),
        }

        try:
            pdf_path = generate_pdf(
                payload,
                template_path=str(TEMPLATES / "report_template.docx"),
                output_dir=str(OUT_DIR),
                base_dir=str(ROOT),
            )
            generated.append(Path(pdf_path).name)
        except Exception:
            # Fallback: create a simple DOCX (or TXT) so the UI still returns a download
            note = "PDF converter not available or file was locked; generated a DOCX/TXT fallback."
            fallback_name = f"{uid}.docx" if DocxDocument else f"{uid}.txt"
            fallback_path = OUT_DIR / fallback_name

            try:
                if DocxDocument:
                    doc = DocxDocument()
                    doc.add_heading("AutoAudit – Finding", 0)
                    doc.add_paragraph(f"Strategy: {strategy.name}")
                    doc.add_paragraph(f"Test ID: {payload['TestID']}")
                    doc.add_paragraph(f"Pass/Fail: {payload['Pass/Fail']}")
                    if payload["Recommendation"]:
                        doc.add_paragraph(f"Recommendation: {payload['Recommendation']}")
                    if payload["Evidence Extract"]:
                        doc.add_paragraph("Evidence:")
                        doc.add_paragraph(payload["Evidence Extract"])
                    doc.save(str(fallback_path))
                else:
                    with open(fallback_path, "w", encoding="utf-8") as f:
                        f.write(f"Strategy: {strategy.name}\n")
                        f.write(f"Test ID: {payload['TestID']}\n")
                        f.write(f"Pass/Fail: {payload['Pass/Fail']}\n")
                        f.write(f"Recommendation: {payload['Recommendation']}\n")
                        f.write(f"Evidence: {payload['Evidence Extract']}\n")
                generated.append(fallback_path.name)
            except Exception:
                # If even the fallback fails, keep going without a report link
                pass

    return {"ok": True, "findings": findings, "reports": generated, "note": note}


@app.get("/reports/{filename}")
def download_report(filename: str):
    path = OUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report not found")

    media = (
        "application/pdf"
        if filename.lower().endswith(".pdf")
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if filename.lower().endswith(".docx")
        else "text/plain"
    )
    return FileResponse(str(path), media_type=media, filename=filename)
