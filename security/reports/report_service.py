"""
report_service.py
=================
Takes a JSON dataset and a Word template, fills every {placeholder} token,
and saves the result as a .docx or PDF.

Quick start::

    from report_service import generate_full_report_docx
    import json

    with open("dataset.json") as f:
        data = json.load(f)

    out = generate_full_report_docx(data)
    # open out in Word, check layout, export to PDF

For headless pipelines use generate_full_report_pdf() instead.

Dataset schema
--------------
The top-level keys the service reads are:

    tenant            dict   org name, domain, assessor, dates, etc.
    summary           dict   scores, risk posture, strengths, top risks
    controls          list   one dict per control (see _single_control_mapping)
    evidence_register list   evidence items for Appendix A (up to 10)
    remediation_plan  list   remediation rows (up to 8)

None of these are required — missing keys produce empty strings in the output.
Key names are normalised before lookup (lower-cased, underscores/hyphens/slashes
collapsed to spaces), so "Tenant_Name", "tenant name", and "tenant-name" all
resolve to the same value.

Adding new tokens
-----------------
1. Add the {New_Token} placeholder to the Word template.
2. Add "New_Token": _pick(n, "new token", "new_token") to the relevant
   mapping function below (_map_tenant, _map_summary, _single_control_mapping,
   etc.).  That's it — no other changes needed.

PDF conversion
--------------
Tries three methods in order:
  1. docx2pdf  (requires Microsoft Word on Windows/macOS)
  2. LibreOffice headless  (soffice must be on PATH)
  3. fpdf2 plain-text fallback  (no layout fidelity, last resort)

If none of those are available the call raises RuntimeError.
"""

from __future__ import annotations

import logging
import os
import re
import subprocess
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Mapping, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

log = logging.getLogger(__name__)

# OOXML tag names we reference directly in lxml operations.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_T  = f"{{{_W_NS}}}t"
_W_R  = f"{{{_W_NS}}}r"
_W_P  = f"{{{_W_NS}}}p"


# ---------------------------------------------------------------------------
# Text sanitisation
#
# Word XML only accepts characters in the Latin-1 range for most text nodes,
# so we swap common Unicode punctuation for ASCII equivalents before writing
# anything back into the document.
# ---------------------------------------------------------------------------

_CHAR_SUBS: Dict[str, str] = {
    "\u2019": "'",  "\u2018": "'",   # smart quotes
    "\u201c": '"',  "\u201d": '"',
    "\u2013": "-",  "\u2014": "-",   # en/em dash
    "\u2022": "*",  "\u2026": "...", # bullet, ellipsis
    "\u00a0": " ",  "\u2192": "->",  # NBSP, right arrow
}


def _sanitise(v: str) -> str:
    for ch, rep in _CHAR_SUBS.items():
        v = v.replace(ch, rep)
    return v


def _sanitise_mapping(m: Dict[str, str]) -> Dict[str, str]:
    return {k: _sanitise(v) for k, v in m.items()}


# ---------------------------------------------------------------------------
# Key normalisation
#
# Dataset keys arrive in all kinds of formats — snake_case, Title Case,
# kebab-case — so we normalise everything to lowercase space-separated words
# before comparing.  _pick() then tries several name variants for each token
# so callers don't need to guess the exact key their dataset uses.
# ---------------------------------------------------------------------------

def _normalize_keys(d: Mapping[str, Any]) -> Dict[str, str]:
    norm: Dict[str, str] = {}
    for k, v in d.items():
        key = " ".join(
            str(k).strip().lower()
            .replace("_", " ").replace("-", " ").replace("/", " ")
            .split()
        )
        norm[key] = "" if v is None else str(v)
    return norm


def _pick(norm: Dict[str, str], *names: str) -> str:
    """Try each name in order and return the first match, or empty string."""
    for n in names:
        key = " ".join(n.strip().lower().split())
        if key in norm:
            return norm[key]
    return ""


# ---------------------------------------------------------------------------
# Placeholder variant expansion
#
# Word occasionally substitutes a Unicode dash for the ASCII hyphen inside
# token names when the template is edited on macOS.  We add aliases for all
# four Unicode dash variants so {Sub-Strategy} and {Sub‑Strategy} both work.
# ---------------------------------------------------------------------------

_UNICODE_DASHES = ["\u2010", "\u2011", "\u2013", "\u2014"]


def _expand_placeholder_variants(mapping: Dict[str, str]) -> None:
    extras: Dict[str, str] = {}
    for k, v in list(mapping.items()):
        if "-" in k:
            for dash in _UNICODE_DASHES:
                alias = k.replace("-", dash)
                if alias not in mapping:
                    extras[alias] = v
    mapping.update(extras)


# ---------------------------------------------------------------------------
# XML substitution engine
#
# Word splits paragraph text across multiple <w:r> runs for formatting
# reasons, so a single token like {Control_Name} can end up split as
# "{Control_", "Name}".  We handle this with two passes:
#
#   Pass 1 — single-run substitution: replaces tokens that happen to fall
#             entirely within one <w:t> node.  Fast and covers most cases.
#
#   Pass 2 — paragraph merge: concatenates all run text, does the replace,
#             then writes the result back into the first <w:t> node and
#             blanks the rest.  This preserves run formatting (bold, colour,
#             font size) because we're reusing existing <w:r> elements rather
#             than creating new ones.
#
# JSON in Extract values contains its own braces, so we escape them to
# sentinel strings before substitution and restore them afterwards.
# ---------------------------------------------------------------------------

def _collapse_token_whitespace(text: str) -> str:
    """Strip whitespace inside {...} so mid-word-wrapped tokens still match."""
    def _strip(m: re.Match) -> str:
        return "{" + re.sub(r"\s+", "", m.group(1)) + "}"
    return re.sub(r"\{([^}]*)\}", _strip, text)


def _sub_mapping_in_element(element, mapping: Dict[str, str]) -> None:
    # Pass 1: single-run tokens
    for t in element.iter():
        if t.tag == _W_T and t.text:
            new = t.text
            for k, v in mapping.items():
                new = new.replace("{" + k + "}", v)
            if new != t.text:
                t.text = new

    # Pass 2: tokens split across runs within a paragraph
    for para in element.iter():
        if para.tag != _W_P:
            continue
        runs = [c for c in para if c.tag == _W_R]
        if not runs:
            continue

        # Collect the first <w:t> from each run (None if the run has no text node)
        t_nodes: List[Optional[Any]] = []
        for r in runs:
            for c in r:
                if c.tag == _W_T:
                    t_nodes.append(c)
                    break
            else:
                t_nodes.append(None)

        full = "".join((t.text or "") for t in t_nodes if t is not None)
        if not full:
            continue

        new = _collapse_token_whitespace(full)
        for k, v in mapping.items():
            new = new.replace("{" + k + "}", v)

        if new != full:
            valid = [t for t in t_nodes if t is not None]
            if valid:
                valid[0].text = new
                for t in valid[1:]:
                    t.text = ""

    # Restore JSON braces that were escaped before substitution
    for t in element.iter():
        if t.tag == _W_T and t.text:
            if "[[LBRACE]]" in t.text or "[[RBRACE]]" in t.text:
                t.text = t.text.replace("[[LBRACE]]", "{").replace("[[RBRACE]]", "}")


def _sub_part(part, mapping: Dict[str, str]) -> None:
    _sub_mapping_in_element(part.element, mapping)


def _iter_all_hf_parts(doc: Document):
    """Yield every header/footer part for all sections and page variants."""
    attrs = (
        "header", "footer",
        "first_page_header", "first_page_footer",
        "even_page_header",  "even_page_footer",
    )
    for section in doc.sections:
        for attr in attrs:
            try:
                hf = getattr(section, attr, None)
                if hf and hf.part:
                    yield hf.part
            except Exception:
                pass


# ---------------------------------------------------------------------------
# Global mapping builders
# ---------------------------------------------------------------------------

def _map_tenant(t: Mapping[str, Any]) -> Dict[str, str]:
    n = _normalize_keys(t)
    return {
        "Tenant_Name":       _pick(n, "tenant name", "tenant_name"),
        "Tenant_Domain":     _pick(n, "tenant domain", "tenant_domain"),
        "Assessor_Name":     _pick(n, "assessor name", "assessor_name"),
        "Frameworks_Used":   _pick(n, "frameworks used", "frameworks_used", "framework"),
        "Assessment_Period": _pick(n, "assessment period", "assessment_period"),
        "Assessment_Date":   _pick(n, "assessment date", "assessment_date"),
        "Classification":    _pick(n, "classification"),
        "Report_Version":    _pick(n, "report version", "report_version", "version"),
        "Distribution":      _pick(n, "distribution"),
        "Prepared_By":       _pick(n, "prepared by", "prepared_by"),
        "Reviewed_By":       _pick(n, "reviewed by", "reviewed_by"),
        "Team_Function":     _pick(n, "team function", "team_function"),
        "Limitations":       _pick(n, "limitations"),
        "Scope_Owner":       _pick(n, "scope owner", "scope_owner"),
    }


def _map_summary(s: Mapping[str, Any]) -> Dict[str, str]:
    n  = _normalize_keys(s)
    rp = _pick(n, "overall risk posture", "overall_risk_posture")
    ex = (
        _pick(n, "executive summary", "executive_summary")
        or _pick(n, "key recommendation", "key_recommendation")
    )
    return {
        "Executive_Summary":      ex,
        "Key_Recommendation":     _pick(n, "key recommendation", "key_recommendation"),
        "Overall_Score":          _pick(n, "overall score", "overall_score"),
        "Overall_Risk_Posture":   rp,
        "OVERALL_RISK_POSTURE":   rp,  # template uses both cases
        "Total_Controls":         _pick(n, "total controls", "total_controls"),
        "Total_Pass":             _pick(n, "total pass", "total_pass"),
        "Total_Fail":             _pick(n, "total fail", "total_fail"),
        "Total_Critical":         _pick(n, "total critical", "total_critical"),
        "Total_High":             _pick(n, "total high", "total_high"),
        "Total_Medium":           _pick(n, "total medium", "total_medium"),
        "Total_Low":              _pick(n, "total low", "total_low"),
        "Top_Risk_1":             _pick(n, "top risk 1", "top_risk_1"),
        "Top_Risk_2":             _pick(n, "top risk 2", "top_risk_2"),
        "Top_Risk_3":             _pick(n, "top risk 3", "top_risk_3"),
        "Strength_1":             _pick(n, "strength 1", "strength_1"),
        "Strength_1_Evidence":    _pick(n, "strength 1 evidence", "strength_1_evidence"),
        "Strength_2":             _pick(n, "strength 2", "strength_2"),
        "Strength_2_Evidence":    _pick(n, "strength 2 evidence", "strength_2_evidence"),
        "Strength_3":             _pick(n, "strength 3", "strength_3"),
        "Strength_3_Evidence":    _pick(n, "strength 3 evidence", "strength_3_evidence"),
        "Strength_4":             _pick(n, "strength 4", "strength_4"),
        "Strength_4_Evidence":    _pick(n, "strength 4 evidence", "strength_4_evidence"),
        "Strength_5":             _pick(n, "strength 5", "strength_5"),
        "Strength_5_Evidence":    _pick(n, "strength 5 evidence", "strength_5_evidence"),
        "Top_Remediation_Action": _pick(n, "top remediation action", "top_remediation_action"),
    }


def _map_categories(data_summary: Mapping[str, Any]) -> Dict[str, str]:
    """
    Handles both dataset shapes:
      - flat:   summary.Cat_1_Pass
      - nested: summary.categories.Cat_1.Pass
    """
    result: Dict[str, str] = {}
    n_summary   = _normalize_keys(data_summary)
    nested_cats = data_summary.get("categories", {})

    for i in range(1, 10):
        flat_pass    = _pick(n_summary, f"cat {i} pass",    f"cat_{i}_pass")
        flat_fail    = _pick(n_summary, f"cat {i} fail",    f"cat_{i}_fail")
        flat_total   = _pick(n_summary, f"cat {i} total",   f"cat_{i}_total")
        flat_comment = _pick(n_summary, f"cat {i} comment", f"cat_{i}_comment")

        if not flat_pass and nested_cats:
            c  = nested_cats.get(f"Cat_{i}", {})
            nc = _normalize_keys(c)
            flat_pass    = _pick(nc, "pass")    or "0"
            flat_fail    = _pick(nc, "fail")    or "0"
            flat_total   = _pick(nc, "total")   or "0"
            flat_comment = _pick(nc, "comment") or ""

        result[f"Cat_{i}_Pass"]    = flat_pass    or "0"
        result[f"Cat_{i}_Fail"]    = flat_fail    or "0"
        result[f"Cat_{i}_Total"]   = flat_total   or "0"
        result[f"Cat_{i}_Comment"] = flat_comment or ""

    return result


def _map_evidence_register(ev: list) -> Dict[str, str]:
    """Build Evidence_N_* keys for Appendix A (up to 10 items)."""
    result: Dict[str, str] = {}
    for i, item in enumerate(ev[:10], 1):
        n = _normalize_keys(item)
        result[f"Evidence_{i}_Description"] = (
            _pick(n, "evidence description", "evidence_description") or _pick(n, "description")
        )
        result[f"Evidence_{i}_Source"] = (
            _pick(n, "evidence source", "evidence_source") or _pick(n, "source")
        )
        result[f"Evidence_{i}_MappedControl"] = (
            _pick(n, "mapped control", "mapped_control") or _pick(n, "uniqueid", "unique id")
        )
        result[f"Evidence_{i}_Date"] = (
            _pick(n, "date captured", "date_captured")
            or _pick(n, "date", "assessment date", "assessment_date")
        )

    # Blank out any template rows that have no corresponding data
    for i in range(len(ev) + 1, 11):
        result.setdefault(f"Evidence_{i}_Description",   "")
        result.setdefault(f"Evidence_{i}_Source",        "")
        result.setdefault(f"Evidence_{i}_MappedControl", "")
        result.setdefault(f"Evidence_{i}_Date",          "")

    return result


def _build_global_mapping(data: Mapping[str, Any]) -> Dict[str, str]:
    """Merge all top-level mappings into one dict for global substitution."""
    m: Dict[str, str] = {}
    m.update(_map_tenant(data.get("tenant", {})))
    m.update(_map_summary(data.get("summary", {})))
    m.update(_map_categories(data.get("summary", {})))
    m.update(_map_evidence_register(data.get("evidence_register", [])))

    now = datetime.now().strftime("%d %b %Y %H:%M")
    # The footer token can appear as '{Date Generated}' or '{Generated Date}'.
    # _collapse_token_whitespace strips spaces inside braces, so both the
    # spaced and spaceless forms need to be registered.
    m["Date Generated"] = now
    m["Generated Date"] = now
    m["Generated_Date"] = now
    m["DateGenerated"]  = now
    m["GeneratedDate"]  = now

    return _sanitise_mapping(m)


# ---------------------------------------------------------------------------
# Per-control mapping
# ---------------------------------------------------------------------------

def _single_control_mapping(ctrl: Mapping[str, Any]) -> Dict[str, str]:
    """
    Build the token map for one finding block.

    Extract values often contain JSON with their own braces.  We escape those
    to sentinel strings here and _sub_mapping_in_element() restores them after
    all {Token} replacements are done.
    """
    n  = _normalize_keys(ctrl)
    ss = _pick(n, "sub strategy", "sub_strategy", "substrategy")

    extract_raw  = _pick(n, "extract", "evidence extract")
    extract_safe = extract_raw.replace("{", "[[LBRACE]]").replace("}", "[[RBRACE]]")

    return {
        "Control Name":         _pick(n, "control name", "control_name"),
        "Control_Name":         _pick(n, "control name", "control_name"),
        "CIS_Section":          _pick(n, "cis section",  "cis_section"),
        "ISO_Mapping":          _pick(n, "iso mapping",  "iso_mapping"),
        "Strategy":             _pick(n, "strategy"),
        "Sub-Strategy":         ss,
        "Sub_Strategy":         ss,
        "Test_id":              _pick(n, "test id", "test_id", "testid"),
        "Level":                _pick(n, "level"),
        "Compliance_Status":    _pick(n, "compliance status",   "compliance_status"),
        "Risk_Rating":          _pick(n, "risk rating",         "risk_rating"),
        "Priority":             _pick(n, "priority"),
        "Pass/Fail":            _pick(n, "pass fail", "pass/fail", "passfail"),
        "Description":          _pick(n, "description"),
        "Observations":         _pick(n, "observations", "observation"),
        "Justification":        _pick(n, "justification"),
        "Evidence_Type":        _pick(n, "evidence type", "evidence_type"),
        "File Name":            _pick(n, "file name", "file_name", "filename"),
        "Extract":              extract_safe,
        "Confidence":           _pick(n, "confidence"),
        "Evidence_Explanation": _pick(n, "evidence explanation", "evidence_explanation"),
        "Impact":               _pick(n, "impact"),
        "Root_Cause":           _pick(n, "root cause", "root_cause"),
        "Remediation":          _pick(n, "remediation"),
        "Recommendations":      _pick(n, "recommendations", "recommendation"),
        "Owner":                _pick(n, "owner"),
        "Target_Date":          _pick(n, "target date", "target_date"),
        "Remediation_Status":   _pick(n, "remediation status", "remediation_status", "status"),
        "UniqueID":             _pick(n, "uniqueid", "unique id", "userid"),
    }


# ---------------------------------------------------------------------------
# Severity bucketing
# ---------------------------------------------------------------------------

def _bucket_fails_by_severity(controls: list) -> Dict[str, Optional[dict]]:
    """
    Pick the first FAIL at each severity level.  The template has one finding
    block per level, so only the first match matters.
    """
    order   = ["Critical", "High", "Medium", "Low"]
    buckets: Dict[str, Optional[dict]] = {s: None for s in order}
    for ctrl in controls:
        n  = _normalize_keys(ctrl)
        pf = _pick(n, "pass fail", "pass/fail", "passfail").upper()
        if not pf:
            status = _pick(n, "compliance status", "compliance_status").lower()
            if "non" in status or "partial" in status:
                pf = "FAIL"
        sv = _pick(n, "risk rating", "risk_rating", "severity")
        if pf == "FAIL" and sv in buckets and buckets[sv] is None:
            buckets[sv] = dict(ctrl)
    return buckets


# ---------------------------------------------------------------------------
# Table location helpers
#
# _find_remediation_table runs before substitution so the Remediation_Action_N
# tokens are still present and searchable.
#
# _find_appendix_b_table and _find_evidence_table also run before substitution,
# identifying tables by their header placeholder text.  The row-level fill
# functions (_substitute_appendix_b, _substitute_evidence_table) use index-
# based row selection so they work correctly even after earlier passes have
# consumed the placeholder content.
# ---------------------------------------------------------------------------

def _find_remediation_table(doc: Document) -> Optional[Any]:
    for table in doc.tables:
        for row in table.rows:
            if re.search(r"Remediation_Action_\d+", " ".join(c.text for c in row.cells)):
                return table
    return None


def _find_appendix_b_table(doc: Document) -> Optional[Any]:
    for table in doc.tables:
        full = " ".join(c.text for row in table.rows for c in row.cells)
        if "Control_Name" in full and "UniqueID" in full:
            return table
    return None


def _find_evidence_table(doc: Document) -> Optional[Any]:
    for table in doc.tables:
        full = " ".join(c.text for row in table.rows for c in row.cells)
        if "Evidence_1_Description" in full:
            return table
    return None


# ---------------------------------------------------------------------------
# Document-level substitution steps
# ---------------------------------------------------------------------------

# Section numbers that anchor each severity finding block in the template.
_BLOCK_SENTINELS = {
    "Critical": "6.1",
    "High":     "6.2",
    "Medium":   "6.3",
    "Low":      "6.4",
}

# These keywords appear in the real section headings but not in TOC entries.
# Requiring both the number AND a keyword prevents matching the TOC.
_BLOCK_SEVERITY_KEYWORDS = {
    "Critical": ["(Critical)", "Critical"],
    "High":     ["(High)",     "High"],
    "Medium":   ["(Medium)",   "Medium"],
    "Low":      ["(Low)",      "Low"],
}


def _substitute_finding_blocks(
    doc: Document,
    severity_controls: Dict[str, Optional[dict]],
) -> None:
    """
    Substitute tokens in each severity block (6.1-6.4), scoped tightly so
    one block's data never leaks into another.

    The TOC contains the same '6.1', '6.2' etc. strings and appears before
    the actual headings in the document body.  We skip TOC entries by
    requiring both the section number and the severity keyword in the same
    element.

    Section 7 is the hard stop — nothing past it gets touched here.
    """
    body     = doc.element.body
    children = list(body)
    order    = ["Critical", "High", "Medium", "Low"]

    anchors: Dict[str, int]     = {}
    section7_idx: Optional[int] = None

    for idx, child in enumerate(children):
        text = "".join(t.text for t in child.iter() if t.tag == _W_T and t.text)

        if (section7_idx is None
                and re.match(r"\s*7[\.\s]", text)
                and "\u2026" not in text
                and "..." not in text
                and len(text) < 80):
            section7_idx = idx

        for sev, sentinel in _BLOCK_SENTINELS.items():
            if sev not in anchors:
                has_num = bool(re.search(r"(?<!\d)" + re.escape(sentinel) + r"(?!\d)", text))
                has_kw  = any(kw in text for kw in _BLOCK_SEVERITY_KEYWORDS[sev])
                if has_num and has_kw:
                    anchors[sev] = idx

    hard_stop = section7_idx if section7_idx is not None else len(children)

    for i, sev in enumerate(order):
        ctrl = severity_controls.get(sev)
        if not ctrl or sev not in anchors:
            continue

        start = anchors[sev]
        end   = hard_stop
        for ns in order[i + 1:]:
            if ns in anchors and anchors[ns] < end:
                end = anchors[ns]

        ctrl_map = _sanitise_mapping(_single_control_mapping(ctrl))
        _expand_placeholder_variants(ctrl_map)
        for child in children[start:end]:
            _sub_mapping_in_element(child, ctrl_map)


def _substitute_remediation_rows(doc: Document, rows: list, rem_table) -> None:
    """Fill the remediation plan table row by row (up to 8 rows)."""
    if not rem_table:
        return

    data_rows = [
        row for row in rem_table.rows
        if re.search(r"Remediation_Action_\d+", " ".join(c.text for c in row.cells))
    ]

    for i, (row, item) in enumerate(zip(data_rows, rows), 1):
        n      = _normalize_keys(item)
        action = (
            _pick(n, "remediation action", "action")
            or _pick(n, f"remediation action {i}", f"remediation_action_{i}")
        )
        row_map = _sanitise_mapping({
            f"Remediation_Action_{i}": action,
            "Owner":              _pick(n, "owner") or "",
            "Target_Date":        _pick(n, "target date", "target_date") or "",
            "Remediation_Status": _pick(n, "status", "remediation status", "remediation_status") or "",
        })
        _sub_mapping_in_element(row._tr, row_map)

    # Blank template rows that have no data
    for i in range(len(rows) + 1, len(data_rows) + 1):
        pad = {
            f"Remediation_Action_{i}": "",
            "Owner": "", "Target_Date": "", "Remediation_Status": "",
        }
        _sub_mapping_in_element(data_rows[i - 1]._tr, pad)


def _substitute_appendix_b(doc: Document, controls: list, app_b_table) -> None:
    """Fill the Appendix B controls summary table row by row."""
    if not app_b_table:
        return

    # Skip header row; fill data rows by index so token replacement order
    # doesn't matter (tokens are already consumed before this runs).
    data_rows = list(app_b_table.rows)[1:]

    for row, ctrl in zip(data_rows, controls):
        n = _normalize_keys(ctrl)
        # Derive Pass/Fail from Compliance_Status if the field is absent.
        pf = _pick(n, "pass fail", "pass/fail")
        if not pf:
            status = _pick(n, "compliance status", "compliance_status").lower()
            pf = "FAIL" if ("non" in status or "partial" in status) else "PASS"
        row_map = _sanitise_mapping({
            "Control_Name": _pick(n, "control name", "control_name") or "",
            "Control Name": _pick(n, "control name", "control_name") or "",
            "UniqueID":     _pick(n, "uniqueid", "unique id") or "",
            "Level":        _pick(n, "level") or "",
            "Pass/Fail":    pf,
        })
        _sub_mapping_in_element(row._tr, row_map)

    for row in data_rows[len(controls):]:
        pad = {"Control_Name": "", "Control Name": "", "UniqueID": "", "Level": "", "Pass/Fail": ""}
        _sub_mapping_in_element(row._tr, pad)


def _substitute_evidence_table(doc: Document, evidence_items: list, ev_table) -> None:
    """
    Fill the Evidence Register table row by row (Appendix A).

    Uses numbered tokens (Evidence_N_MappedControl, Evidence_N_Date) rather
    than shared token names (UniqueID, Assessment_Date) to avoid cross-
    contamination with the global substitution pass.

    Rows are selected by index rather than by token text, since this function
    runs after earlier passes have already consumed the placeholder content.
    """
    if not ev_table:
        return

    # Skip header row; fill data rows by index.
    data_rows = list(ev_table.rows)[1:]

    for i, (row, item) in enumerate(zip(data_rows, evidence_items), 1):
        n      = _normalize_keys(item)
        desc   = _pick(n, "evidence description", "evidence_description") or _pick(n, "description") or ""
        source = _pick(n, "evidence source", "evidence_source") or _pick(n, "source") or ""
        mapped = _pick(n, "mapped control", "mapped_control") or _pick(n, "uniqueid", "unique id") or ""
        date   = (
            _pick(n, "date captured", "date_captured")
            or _pick(n, "date", "assessment date", "assessment_date") or ""
        )
        row_map = _sanitise_mapping({
            f"Evidence_{i}_Description":   desc,
            f"Evidence_{i}_Source":        source,
            f"Evidence_{i}_MappedControl": mapped,
            f"Evidence_{i}_Date":          date,
        })
        _sub_mapping_in_element(row._tr, row_map)

    for idx in range(len(evidence_items) + 1, len(data_rows) + 1):
        pad = {
            f"Evidence_{idx}_Description":   "",
            f"Evidence_{idx}_Source":        "",
            f"Evidence_{idx}_MappedControl": "",
            f"Evidence_{idx}_Date":          "",
        }
        _sub_mapping_in_element(data_rows[idx - 1]._tr, pad)


def _substitute_global(doc: Document, mapping: Dict[str, str]) -> None:
    """Global pass — runs after all scoped steps, fills everything remaining."""
    _sub_part(doc.part, mapping)
    for part in _iter_all_hf_parts(doc):
        _sub_part(part, mapping)


# ---------------------------------------------------------------------------
# Stray marker removal
# ---------------------------------------------------------------------------

_STRAY_MARKERS = [
    "[Embed evidence here]",
    "[Embed screenshot here]",
    "End of report. Delete unused finding blocks and this note before issuing the final version.",
    "[Repeat the finding block above for each additional failing control, ordered Critical \u2192 High \u2192 Medium \u2192 Low]",
    "[Repeat the finding block above for each additional failing control, ordered Critical -> High -> Medium -> Low]",
]


def _remove_markers(doc: Document, extra_markers: Optional[List[str]] = None) -> None:
    """Strip template instruction strings that should not appear in output."""
    markers = _STRAY_MARKERS + (extra_markers or [])

    def _scrub(element) -> None:
        for t in element.iter():
            if t.tag == _W_T and t.text:
                new = t.text
                for m in markers:
                    new = new.replace(m, "")
                if new != t.text:
                    t.text = new

    _scrub(doc.element)
    for part in _iter_all_hf_parts(doc):
        _scrub(part.element)


# ---------------------------------------------------------------------------
# Image embedding
# ---------------------------------------------------------------------------

def _insert_image_at_marker(
    doc: Document,
    marker: str,
    image_path: os.PathLike | str,
    width_inches: float = 6.0,
) -> bool:
    """Replace marker text with an embedded image. Returns False if not found."""
    ip = Path(image_path)
    if not ip.exists():
        return False

    target = None
    for p in doc.paragraphs:
        if marker in "".join(r.text for r in p.runs):
            target = p
            break

    if not target:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if marker in "".join(r.text for r in p.runs):
                            target = p
                            break

    if not target:
        return False

    for r in target.runs:
        r.text = r.text.replace(marker, "")
    target.add_run().add_picture(str(ip), width=Inches(width_inches))
    return True


# ---------------------------------------------------------------------------
# Finding table column width fix
#
# The template's finding tables (6.1-6.4) have correct column widths in their
# XML (tblW, tcW) but are missing the tblLayout element.  Without it the OOXML
# spec defaults to autofit, which lets Word override those widths when long
# content is injected — e.g. a JSON extract in the Justification row collapses
# the right column.  Adding tblLayout type="fixed" is all that's needed.
#
# We identify finding tables by checking whether at least 4 known row labels
# appear in the left column, so the fix is robust to table index shifts if the
# template is ever restructured.
# ---------------------------------------------------------------------------

_FINDING_ROW_LABELS = {
    "Description", "Observation", "Justification",
    "Evidence Reviewed", "Evidence Explanation",
    "Risk / Impact", "Root Cause", "Recommendation",
    "Control Reference", "Strategy / Sub-strategy", "Test ID",
    "CIS Level", "Compliance Status", "Risk Rating",
    "Priority", "Result", "Owner",
}


def _fix_finding_table_widths(doc: Document) -> None:
    for table in doc.tables:
        left_labels = {row.cells[0].text.strip() for row in table.rows if row.cells}
        if len(left_labels & _FINDING_ROW_LABELS) < 4:
            continue
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        if tblPr.find(qn("w:tblLayout")) is None:
            el = OxmlElement("w:tblLayout")
            el.set(qn("w:type"), "fixed")
            tblPr.append(el)


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------

def _convert_docx_to_pdf(input_docx: Path, output_pdf: Path) -> None:
    """
    Try docx2pdf, then LibreOffice, then an fpdf2 plain-text fallback.
    Raises RuntimeError if all three fail.
    """
    # docx2pdf (needs Microsoft Word)
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(input_docx), str(output_pdf))
        log.info("PDF via docx2pdf: %s", output_pdf)
        return
    except ImportError:
        log.debug("docx2pdf not installed, trying LibreOffice.")
    except Exception as exc:
        log.warning("docx2pdf failed (%s), trying LibreOffice.", exc)

    # LibreOffice headless
    try:
        out_dir = str(output_pdf.parent.resolve())
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, str(input_docx.resolve())],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        lo_out = output_pdf.parent / (input_docx.stem + ".pdf")
        if lo_out.exists() and lo_out != output_pdf:
            lo_out.replace(output_pdf)
        log.info("PDF via LibreOffice: %s", output_pdf)
        return
    except FileNotFoundError:
        log.debug("soffice not found, using plain-text fallback.")
    except subprocess.CalledProcessError as exc:
        log.warning("LibreOffice failed (%s), using plain-text fallback.", exc)

    # fpdf2 plain-text fallback (no layout fidelity)
    if _simple_pdf_from_docx(input_docx, output_pdf):
        log.warning("PDF via fpdf2 fallback (no formatting): %s", output_pdf)
        return

    raise RuntimeError(
        "PDF conversion failed — install docx2pdf (Word) or LibreOffice, "
        "or ensure fpdf2 is available."
    )


def _simple_pdf_from_docx(input_docx: Path, output_pdf: Path) -> bool:
    """Text-only PDF via fpdf2. Used only when Word and LibreOffice are unavailable."""
    try:
        from fpdf import FPDF  # type: ignore

        doc = Document(str(input_docx))
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("helvetica", size=12)

        def _safe(text: str) -> str:
            if not text:
                return ""
            try:
                return text.encode("latin-1", errors="replace").decode("latin-1")
            except Exception:
                return text.encode("ascii", errors="replace").decode("ascii")

        def _line(text: str) -> None:
            text = _safe(text.strip())
            if text:
                pdf.multi_cell(0, 8, text)
                pdf.ln(1)

        for p in doc.paragraphs:
            _line(p.text or "")
        for tbl in doc.tables:
            for row in tbl.rows:
                _line(" | ".join((c.text or "").strip() for c in row.cells))

        pdf.output(str(output_pdf))
        return True
    except Exception as exc:
        log.error("fpdf2 fallback failed: %s", exc)
        return False


# ---------------------------------------------------------------------------
# Render core
# ---------------------------------------------------------------------------

def _resolve_output_stem(mapping: Dict[str, str], override: Optional[str]) -> str:
    if override:
        return override
    tenant = mapping.get("Tenant_Name", "report").replace(" ", "_")
    date   = mapping.get("Assessment_Date", datetime.now().strftime("%d%b%Y")).replace(" ", "")
    return f"{tenant}_{date}_AutoAudit_Report"


def _render_report_doc(
    data: Mapping[str, Any],
    template_path: Path,
) -> Tuple[Document, Dict[str, str]]:
    """
    Load the template and run all substitution passes in the correct order.

    Order matters — do not rearrange:
      1. Locate scoped tables while placeholder text is still raw.
      2. Per-block finding substitution (scoped to sections 6.1–6.4).
      3. Per-row remediation table.
      4. Per-row Appendix B (index-based; Pass/Fail derived from Compliance_Status).
      5. Per-row Evidence Register (index-based; numbered tokens throughout).
      6. Global substitution (fills everything remaining).
      7. Remove stray template instruction markers.
      8. Lock finding table column widths.
    """
    global_mapping = _build_global_mapping(data)
    _expand_placeholder_variants(global_mapping)

    severity_controls = _bucket_fails_by_severity(data.get("controls", []))
    remediation_rows  = data.get("remediation_plan", [])[:8]
    all_controls      = data.get("controls", [])

    doc = Document(str(template_path))

    rem_table   = _find_remediation_table(doc)
    app_b_table = _find_appendix_b_table(doc)
    ev_table    = _find_evidence_table(doc)

    _substitute_finding_blocks(doc, severity_controls)
    _substitute_remediation_rows(doc, remediation_rows, rem_table)
    _substitute_appendix_b(doc, all_controls, app_b_table)
    _substitute_evidence_table(doc, data.get("evidence_register", [])[:10], ev_table)
    _substitute_global(doc, global_mapping)
    _remove_markers(doc)
    _fix_finding_table_widths(doc)

    return doc, global_mapping


# ---------------------------------------------------------------------------
# Public API — full report
# ---------------------------------------------------------------------------

def generate_full_report_docx(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "AutoAudit_Report_Template.docx",
    output_dir: os.PathLike | str = "reports_out",
    output_filename: Optional[str] = None,
) -> Path:
    """
    Fill the template with *data* and save as a .docx.

    The output filename defaults to ``<Tenant>_<Date>_AutoAudit_Report.docx``.
    Open the result in Word to check layout before exporting to PDF, or call
    generate_full_report_pdf() if you want the PDF directly.

    Returns the path to the generated file.
    """
    tpath = Path(template_path)
    if not tpath.exists():
        raise FileNotFoundError(f"Template not found: {tpath}")

    outdir = Path(output_dir)
    outdir.mkdir(parents=True, exist_ok=True)

    doc, mapping = _render_report_doc(data, tpath)
    out = outdir / f"{_resolve_output_stem(mapping, output_filename)}.docx"
    doc.save(str(out))

    log.info("Report written to: %s", out)
    return out


def generate_full_report_pdf(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "AutoAudit_Report_Template.docx",
    output_dir: os.PathLike | str = "reports_out",
    output_filename: Optional[str] = None,
    keep_docx: bool = False,
) -> Path:
    """
    Fill the template, convert to PDF, and return the PDF path.

    Set keep_docx=True to also keep the intermediate .docx alongside the PDF.
    PDF conversion tries docx2pdf → LibreOffice → fpdf2 in that order.

    Returns the path to the generated PDF.
    """
    tpath = Path(template_path)
    if not tpath.exists():
        raise FileNotFoundError(f"Template not found: {tpath}")

    outdir = Path(output_dir)
    outdir.mkdir(parents=True, exist_ok=True)

    doc, mapping  = _render_report_doc(data, tpath)
    stem          = _resolve_output_stem(mapping, output_filename)
    docx_path     = outdir / f"{stem}.docx"
    pdf_path      = outdir / f"{stem}.pdf"

    doc.save(str(docx_path))
    _convert_docx_to_pdf(docx_path, pdf_path)

    if not keep_docx:
        try:
            docx_path.unlink()
        except Exception:
            pass

    return pdf_path


def convert_docx_to_pdf(
    docx_path: os.PathLike | str,
    output_dir: Optional[os.PathLike | str] = None,
) -> Path:
    """
    Convert an existing .docx to PDF.

    Handy after generate_full_report_docx() has been reviewed in Word and you
    want to convert without re-rendering the whole report.

    Returns the path to the generated PDF.
    """
    src = Path(docx_path)
    if not src.exists():
        raise FileNotFoundError(f"File not found: {src}")

    outdir   = Path(output_dir) if output_dir else src.parent
    outdir.mkdir(parents=True, exist_ok=True)
    pdf_path = outdir / (src.stem + ".pdf")

    _convert_docx_to_pdf(src, pdf_path)
    return pdf_path


# ---------------------------------------------------------------------------
# Public API — single finding (legacy)
# ---------------------------------------------------------------------------

def _build_single_finding_mapping(
    data: Mapping[str, Any],
    base_dir: Path,
) -> Tuple[Dict[str, str], Optional[Path], str]:
    n = _normalize_keys(data)

    unique_id = _pick(n, "uniqueid", "unique id", "userid") or str(uuid.uuid4())
    ss        = _pick(n, "sub strategy", "sub-strategy")
    extract   = _pick(n, "evidence extract", "extract")

    ev_str      = _pick(n, "evidence", "evidence path", "file", "file path",
                         "filepath", "image", "screenshot")
    preview_str = _pick(n, "evidence preview", "preview", "embed path")

    file_name:  str            = ""
    embed_path: Optional[Path] = None

    if ev_str:
        ep = Path(ev_str)
        if not ep.is_absolute():
            ep = base_dir / ep
        file_name = ep.name

    if preview_str:
        pp = Path(preview_str)
        if not pp.is_absolute():
            pp = base_dir / pp
        if pp.exists():
            embed_path = pp
    elif ev_str:
        ep = Path(ev_str)
        if not ep.is_absolute():
            ep = base_dir / ep
        if ep.exists() and ep.suffix.lower() in {".png", ".jpg", ".jpeg",
                                                   ".tif", ".tiff", ".bmp", ".webp"}:
            embed_path = ep

    mapping: Dict[str, str] = {
        "UniqueID":        unique_id,
        "Strategy":        _pick(n, "strategy"),
        "Test_id":         _pick(n, "testid", "test id"),
        "Sub-Strategy":    ss,
        "Sub_Strategy":    ss,
        "Level":           _pick(n, "ml level", "level"),
        "Pass/Fail":       _pick(n, "pass fail", "pass/fail"),
        "Priority":        _pick(n, "priority"),
        "Recommendations": _pick(n, "recommendation", "recommendations"),
        "Extract":         extract.replace("{", "[[LBRACE]]").replace("}", "[[RBRACE]]"),
        "Description":     _pick(n, "description"),
        "Confidence":      _pick(n, "confidence"),
        "File Name":       file_name,
        "Date Generated":  datetime.now().strftime("%d %b %Y"),
        "Generated Date":  datetime.now().strftime("%d %b %Y"),
    }
    return _sanitise_mapping(mapping), embed_path, unique_id


def _render_single_finding_doc(
    data: Mapping[str, Any],
    base_dir: Path,
    template_path: Path,
    image_marker: str,
    unique_id_override: Optional[str],
) -> Tuple[Document, str]:
    mapping, embed_path, unique_id = _build_single_finding_mapping(data, base_dir)

    if unique_id_override:
        unique_id           = unique_id_override
        mapping["UniqueID"] = unique_id

    doc = Document(str(template_path))
    _expand_placeholder_variants(mapping)
    _sub_mapping_in_element(doc.element, mapping)
    for part in _iter_all_hf_parts(doc):
        _sub_mapping_in_element(part.element, mapping)

    if embed_path:
        if not _insert_image_at_marker(doc, image_marker, embed_path):
            _insert_image_at_marker(doc, "[Embed screenshot here]", embed_path)
    else:
        _remove_markers(doc, [image_marker, "[Embed screenshot here]"])

    return doc, unique_id


def generate_single_finding_docx(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "AutoAudit_Report_Template.docx",
    output_dir: os.PathLike | str = "reports_out",
    base_dir: os.PathLike | str = ".",
    image_marker: str = "[Embed evidence here]",
    unique_id_override: Optional[str] = None,
) -> Path:
    """
    Render a single finding entry as a .docx (legacy / OCR pipeline usage).
    Returns the path to the generated file.
    """
    tpath = Path(template_path)
    if not tpath.exists():
        raise FileNotFoundError(f"Template not found: {tpath}")

    outdir = Path(output_dir)
    outdir.mkdir(parents=True, exist_ok=True)

    doc, uid = _render_single_finding_doc(
        data, Path(base_dir), tpath, image_marker, unique_id_override
    )
    out = outdir / f"{uid}.docx"
    doc.save(str(out))
    log.info("Single-finding docx: %s", out)
    return out


def generate_single_finding_pdf(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "AutoAudit_Report_Template.docx",
    output_dir: os.PathLike | str = "reports_out",
    base_dir: os.PathLike | str = ".",
    image_marker: str = "[Embed evidence here]",
    unique_id_override: Optional[str] = None,
) -> Path:
    """
    Render a single finding entry directly to PDF (legacy usage).
    Prefer generate_single_finding_docx() when you need to review the output first.
    Returns the path to the generated PDF.
    """
    tpath = Path(template_path)
    if not tpath.exists():
        raise FileNotFoundError(f"Template not found: {tpath}")

    outdir = Path(output_dir)
    outdir.mkdir(parents=True, exist_ok=True)

    doc, uid = _render_single_finding_doc(
        data, Path(base_dir), tpath, image_marker, unique_id_override
    )
    docx_path = outdir / f"{uid}.docx"
    pdf_path  = outdir / f"{uid}.pdf"
    doc.save(str(docx_path))
    _convert_docx_to_pdf(docx_path, pdf_path)

    try:
        docx_path.unlink()
    except Exception:
        pass

    return pdf_path


# ---------------------------------------------------------------------------
# Backward-compatibility alias
#
# security/evidence_backend/reportgenerator.py and security/evidence_ui/app.py
# both import generate_pdf from this module.  It maps to generate_single_finding_pdf
# which accepts the same arguments those callers pass.
# ---------------------------------------------------------------------------

def generate_pdf(
    data: Mapping[str, Any],
    *,
    template_path: os.PathLike | str = "AutoAudit_Report_Template.docx",
    output_dir: os.PathLike | str = "reports_out",
    base_dir: os.PathLike | str = ".",
    image_marker: str = "[Embed evidence here]",
    unique_id_override: Optional[str] = None,
) -> Path:
    """Alias for generate_single_finding_pdf — kept for backward compatibility."""
    return generate_single_finding_pdf(
        data,
        template_path=template_path,
        output_dir=output_dir,
        base_dir=base_dir,
        image_marker=image_marker,
        unique_id_override=unique_id_override,
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")

    args = sys.argv[1:]

    # Subcommand: convert an existing docx to PDF
    if args and args[0] == "convert":
        src    = args[1] if len(args) > 1 else None
        outdir = args[2] if len(args) > 2 else None
        if not src:
            print("Usage: python report_service.py convert <report.docx> [output_dir]")
            sys.exit(1)
        pdf = convert_docx_to_pdf(src, outdir)
        print(f"PDF written to: {pdf}")
        sys.exit(0)

    # Default: generate full report
    dataset  = args[0] if args and not args[0].startswith("--") else "fake_dataset.json"
    template = args[1] if len(args) > 1 and not args[1].startswith("--") else "AutoAudit_Report_Template.docx"
    outdir   = args[2] if len(args) > 2 and not args[2].startswith("--") else "reports_out"
    to_pdf   = "--pdf" in args
    keep     = "--keep-docx" in args

    with open(dataset) as f:
        data = json.load(f)

    if to_pdf:
        out = generate_full_report_pdf(data, template_path=template, output_dir=outdir, keep_docx=keep)
        print(f"PDF written to: {out}")
    else:
        out = generate_full_report_docx(data, template_path=template, output_dir=outdir)
        print(f"Docx written to: {out}")
        print(f"Convert to PDF:  python {sys.argv[0]} convert \"{out}\"")