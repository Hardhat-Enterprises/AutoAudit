import inspect
import logging
import sys
import tempfile
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Tuple


def _locate_security_dir() -> Path:
    base = Path(__file__).resolve()
    for parent in base.parents:
        candidate = parent / "security"
        if candidate.exists():
            return candidate
    raise RuntimeError("Security toolkit directory not found relative to backend-api")


SECURITY_DIR = _locate_security_dir()
ROOT_DIR = SECURITY_DIR.parent

if str(SECURITY_DIR) not in sys.path:
    sys.path.append(str(SECURITY_DIR))

from strategies import load_strategies  # type: ignore  # noqa: E402
from reports.report_service import generate_pdf  # type: ignore  # noqa: E402
from backend.core_ocr import extract_text_and_preview  # type: ignore  # noqa: E402

try:  # Optional dependency for DOCX fallback
    from docx import Document as DocxDocument  # type: ignore
except Exception:  # pragma: no cover - optional
    DocxDocument = None


RESULTS_DIR = SECURITY_DIR / "results"
REPORTS_DIR = RESULTS_DIR / "reports"
PREVIEWS_DIR = RESULTS_DIR / "previews"
UPLOADS_DIR = RESULTS_DIR / "uploads"
TEMPLATE_PATH = RESULTS_DIR / "report_template.docx"

for folder in (RESULTS_DIR, REPORTS_DIR, PREVIEWS_DIR, UPLOADS_DIR):
    folder.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger("evidence")


def list_strategy_options() -> List[Dict[str, str]]:
    """Return strategy names and descriptions from the security toolkit."""
    strategies = []
    for strategy in load_strategies() or []:
        try:
            description = strategy.description()
        except Exception:
            description = ""
        strategies.append({"name": strategy.name, "description": description})
    return strategies


def _safe_uid(value: str) -> str:
    keep = []
    for ch in value:
        if ch.isalnum() or ch in ("-", "_"):
            keep.append(ch)
        else:
            keep.append("-")
    cleaned = "".join(keep).strip("-")
    return cleaned or "evidence"


def _save_upload(filename: str, data: bytes) -> Path:
    suffix = Path(filename or "evidence").suffix or ".bin"
    stem = _safe_uid(Path(filename or "evidence").stem)
    upload_path = UPLOADS_DIR / f"{int(time.time() * 1000)}-{stem}{suffix}"
    upload_path.write_bytes(data)
    return upload_path


def _call_strategy_emit(
    strategy: Any,
    text: str,
    *,
    source_file: str,
    full_path: str | None = None,
    preview_path: str | None = None,
) -> List[Dict[str, Any]] | None:
    emit = getattr(strategy, "emit_hits", None)
    if not emit:
        return None

    try:
        sig = inspect.signature(emit)
        params = sig.parameters
        accepts_kwargs = any(p.kind == inspect.Parameter.VAR_KEYWORD for p in params.values())
    except (TypeError, ValueError):
        params = {}
        accepts_kwargs = True

    def _include(name: str) -> bool:
        return accepts_kwargs or name in params

    kwargs: Dict[str, Any] = {}
    if _include("source_file"):
        kwargs["source_file"] = source_file
    if full_path and _include("full_path"):
        kwargs["full_path"] = full_path
    if preview_path and _include("preview_path"):
        kwargs["preview_path"] = preview_path

    return emit(text, **kwargs) or []


def _evaluate_strategy_text(
    strategy: Any,
    text: str,
    *,
    evidence_name: str,
    full_path: Path,
    preview_path: Path | None,
) -> List[Dict[str, Any]]:
    preview_str = str(preview_path) if preview_path else ""
    findings: List[Dict[str, Any]] = []

    emitted = _call_strategy_emit(
        strategy,
        text,
        source_file=evidence_name,
        full_path=str(full_path),
        preview_path=preview_str,
    )

    if emitted is not None:
        for raw in emitted:
            if not isinstance(raw, dict):
                continue
            row = dict(raw)
            row.setdefault("evidence_file", evidence_name)
            if preview_str and not row.get("preview_path"):
                row["preview_path"] = preview_str
            findings.append(row)
        return findings

    hits = strategy.match(text) or []
    if not hits:
        return []

    row = {
        "test_id": "",
        "sub_strategy": "",
        "detected_level": "",
        "pass_fail": "",
        "priority": "",
        "recommendation": "",
        "evidence": hits,
        "evidence_file": evidence_name,
    }
    if preview_str:
        row["preview_path"] = preview_str
    findings.append(row)
    return findings


def _generate_report(payload: Dict[str, Any]) -> str | None:
    try:
        pdf_path = generate_pdf(
            payload,
            template_path=str(TEMPLATE_PATH),
            output_dir=str(REPORTS_DIR),
            base_dir=str(SECURITY_DIR),
        )
        return Path(pdf_path).name
    except Exception as exc:  # pragma: no cover - fallback path
        logger.warning("PDF generation failed, attempting fallback: %s", exc)
        fallback_name = f"{payload['UniqueID']}.docx" if DocxDocument else f"{payload['UniqueID']}.txt"
        fallback_path = REPORTS_DIR / fallback_name
        try:
            if DocxDocument:
                doc = DocxDocument()
                doc.add_heading("AutoAudit – Finding", level=0)
                doc.add_paragraph(f"Strategy: {payload['Strategy']}")
                if payload.get("TestID"):
                    doc.add_paragraph(f"Test ID: {payload['TestID']}")
                if payload.get("Pass/Fail"):
                    doc.add_paragraph(f"Pass/Fail: {payload['Pass/Fail']}")
                if payload.get("Recommendation"):
                    doc.add_paragraph(f"Recommendation: {payload['Recommendation']}")
                if payload.get("Evidence Extract"):
                    doc.add_paragraph("Evidence Extract:")
                    doc.add_paragraph(payload["Evidence Extract"])
                doc.save(str(fallback_path))
            else:
                with open(fallback_path, "w", encoding="utf-8") as handle:
                    handle.write(f"Strategy: {payload['Strategy']}\n")
                    handle.write(f"Test ID: {payload.get('TestID', '')}\n")
                    handle.write(f"Pass/Fail: {payload.get('Pass/Fail', '')}\n")
                    handle.write(f"Recommendation: {payload.get('Recommendation', '')}\n")
                    handle.write(f"Evidence Extract: {payload.get('Evidence Extract', '')}\n")
            return fallback_path.name
        except Exception as fallback_error:
            logger.error("Unable to create fallback report: %s", fallback_error)
            return None


def scan_evidence(*, filename: str, data: bytes, strategy_name: str, user_id: str = "user") -> Dict[str, Any]:
    strategies = load_strategies() or []
    strategy = next((s for s in strategies if s.name == strategy_name), None)
    if not strategy:
        raise ValueError(f"Unknown strategy: {strategy_name}")

    if not data:
        raise ValueError("Empty evidence upload")

    upload_path = _save_upload(filename, data)
    zip_workspace: tempfile.TemporaryDirectory[str] | None = None
    findings: List[Dict[str, Any]] = []
    text_found = False

    try:
        entries: List[Tuple[Path, str]] = []
        display_name = Path(filename or "evidence").name or "evidence"

        if upload_path.suffix.lower() == ".zip":
            zip_workspace = tempfile.TemporaryDirectory()
            try:
                with zipfile.ZipFile(upload_path) as archive:
                    archive.extractall(zip_workspace.name)
            except zipfile.BadZipFile as exc:
                raise ValueError("Invalid ZIP archive uploaded") from exc

            extracted_root = Path(zip_workspace.name)
            for inner in sorted(p for p in extracted_root.rglob("*") if p.is_file()):
                rel = inner.relative_to(extracted_root).as_posix()
                rel_name = f"{display_name}:{rel}" if rel else display_name
                entries.append((inner, rel_name))
        else:
            entries.append((upload_path, display_name))

        for entry_path, entry_label in entries:
            entry_text, preview_path = extract_text_and_preview(entry_path, PREVIEWS_DIR)
            if not entry_text.strip():
                continue
            text_found = True
            findings.extend(
                _evaluate_strategy_text(
                    strategy,
                    entry_text,
                    evidence_name=entry_label,
                    full_path=entry_path,
                    preview_path=preview_path,
                )
            )

    finally:
        try:
            upload_path.unlink(missing_ok=True)
        except Exception:
            pass
        if zip_workspace:
            zip_workspace.cleanup()

    if not text_found:
        return {
            "strategy": strategy.name,
            "file": filename,
            "findings": [],
            "reports": [],
            "note": "No readable text found in evidence.",
        }

    generated_reports: List[str] = []
    note = ""
    for idx, finding in enumerate(findings, start=1):
        report_id = _safe_uid(f"{user_id}-{strategy.name}-{Path(filename or 'evidence').stem}-{idx}-{uuid.uuid4().hex[:4]}")
        evidence_name = finding.get("evidence_file", filename)
        preview_hint = finding.get("preview_path", "")
        payload = {
            "UniqueID": report_id,
            "UserID": user_id,
            "Evidence": evidence_name,
            "Evidence Preview": preview_hint,
            "Strategy": strategy.name,
            "TestID": finding.get("test_id", ""),
            "Sub-Strategy": finding.get("sub_strategy", ""),
            "ML Level": finding.get("detected_level", ""),
            "Pass/Fail": finding.get("pass_fail", ""),
            "Priority": finding.get("priority", ""),
            "Recommendation": finding.get("recommendation", ""),
            "Evidence Extract": "; ".join(finding.get("evidence", [])),
            "Description": finding.get("description", ""),
            "Confidence": finding.get("confidence", ""),
        }
        report_name = _generate_report(payload)
        if report_name:
            generated_reports.append(report_name)
        else:
            note = "Report generation failed for one or more findings."

    return {
        "strategy": strategy.name,
        "file": filename,
        "findings": findings,
        "reports": generated_reports,
        "note": note,
    }


def resolve_report_path(filename: str) -> Path:
    safe_name = Path(filename).name
    return REPORTS_DIR / safe_name
